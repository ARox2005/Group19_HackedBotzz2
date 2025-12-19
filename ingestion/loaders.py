"""
Document loaders for extracting text from various file formats.
Supports PDF, DOCX, and TXT files.
"""

import os
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


class BaseLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, file_path: str) -> str:
        """Load and extract text from file."""
        pass
    
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        pass


class PDFLoader(BaseLoader):
    """
    PDF document loader using PyMuPDF.
    Extracts text from all pages of a PDF file.
    Falls back to OCR for scanned/image-based PDFs.
    """
    
    def __init__(self, ocr_fallback: bool = True):
        """
        Initialize PDF loader.
        
        Args:
            ocr_fallback: Whether to use OCR for pages with no text (default: True)
        """
        self.ocr_fallback = ocr_fallback
        self._ocr_processor = None
    
    @property
    def ocr_processor(self):
        """Lazy load OCR processor."""
        if self._ocr_processor is None:
            from .ocr import OCRProcessor
            self._ocr_processor = OCRProcessor()
        return self._ocr_processor
    
    def _extract_images_from_page(self, page) -> list:
        """Extract images from a PDF page."""
        images = []
        try:
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = page.parent.extract_image(xref)
                    if base_image:
                        images.append(base_image["image"])
                except Exception:
                    continue
        except Exception:
            pass
        return images
    
    def _ocr_page_as_image(self, page) -> str:
        """Render PDF page as image and OCR it."""
        try:
            # Render page to image (higher resolution for better OCR)
            mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # OCR the rendered page
            text = self.ocr_processor.extract_text_from_bytes(img_bytes, preprocess=True)
            return text
        except Exception as e:
            print(f"OCR failed for page: {e}")
            return ""
    
    def load(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        Falls back to OCR for scanned/image-based PDFs.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text from all pages
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        text_parts = []
        
        try:
            # Open PDF document
            doc = fitz.open(file_path)
            total_pages = len(doc)
            print(f"[PDF Loader] Processing {total_pages} pages from: {file_path}")
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # First try to extract embedded text
                page_text = page.get_text("text")
                
                if page_text.strip():
                    print(f"[PDF Loader] Page {page_num + 1}: Found {len(page_text)} chars of text")
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                elif self.ocr_fallback:
                    # No text found - try OCR
                    print(f"[PDF Loader] Page {page_num + 1}: No text, attempting OCR...")
                    
                    ocr_text = ""
                    
                    # Method 1: Try to extract and OCR embedded images
                    try:
                        images = self._extract_images_from_page(page)
                        if images:
                            print(f"[PDF Loader] Page {page_num + 1}: Found {len(images)} embedded images")
                            for img_bytes in images:
                                try:
                                    img_text = self.ocr_processor.extract_text_from_bytes(img_bytes, preprocess=False)
                                    if img_text.strip():
                                        ocr_text += img_text + "\n"
                                except Exception as img_err:
                                    print(f"[PDF Loader] Image OCR error: {img_err}")
                    except Exception as e:
                        print(f"[PDF Loader] Error extracting images: {e}")
                    
                    # Method 2: If no text from images, render page and OCR
                    if not ocr_text.strip():
                        print(f"[PDF Loader] Page {page_num + 1}: Rendering page as image for OCR...")
                        try:
                            ocr_text = self._ocr_page_as_image(page)
                        except Exception as e:
                            print(f"[PDF Loader] Page render OCR error: {e}")
                    
                    if ocr_text.strip():
                        print(f"[PDF Loader] Page {page_num + 1}: OCR extracted {len(ocr_text)} chars")
                        text_parts.append(f"[Page {page_num + 1} - OCR]\n{ocr_text}")
                    else:
                        print(f"[PDF Loader] Page {page_num + 1}: OCR failed to extract any text")
            
            doc.close()
            
            print(f"[PDF Loader] Total: extracted text from {len(text_parts)} pages")
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise RuntimeError(f"Error reading PDF file: {e}")
        
        return "\n\n".join(text_parts)
    
    def supported_extensions(self) -> List[str]:
        return ['.pdf']


class DOCXLoader(BaseLoader):
    """
    DOCX document loader using python-docx.
    Extracts text from paragraphs and tables.
    """
    
    def load(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from document
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {e}")
        
        return "\n\n".join(text_parts)
    
    def supported_extensions(self) -> List[str]:
        return ['.docx']


class TXTLoader(BaseLoader):
    """
    Plain text file loader.
    Handles various text encodings.
    """
    
    def load(self, file_path: str) -> str:
        """
        Read text from plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            File contents as string
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Text file not found: {file_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise RuntimeError(f"Could not decode text file with any supported encoding")
    
    def supported_extensions(self) -> List[str]:
        return ['.txt']


class DocumentLoader:
    """
    Unified document loader that automatically selects the appropriate loader
    based on file extension.
    """
    
    def __init__(self):
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DOCXLoader(),
            '.txt': TXTLoader(),
        }
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document and return extracted text with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - source: Source file path
                - filename: Base filename
                - extension: File extension
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.loaders:
            raise ValueError(f"Unsupported file format: {extension}")
        
        loader = self.loaders[extension]
        text = loader.load(file_path)
        
        return {
            'text': text,
            'source': str(file_path),
            'filename': path.name,
            'extension': extension
        }
    
    def load_multiple(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Load multiple documents.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of document dictionaries
        """
        documents = []
        for file_path in file_paths:
            try:
                doc = self.load(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"Warning: Failed to load {file_path}: {e}")
        
        return documents
    
    def supported_extensions(self) -> List[str]:
        """Return all supported file extensions."""
        return list(self.loaders.keys())
